"""
ingest.py — Multimodal file ingestion pipeline
Handles: PDF, DOCX, Images (OCR), Audio (Whisper STT)
"""

import os
import fitz                          # PyMuPDF
import docx
import whisper
import pytesseract
from PIL import Image
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ── Whisper model (loaded once, reused) ──────────────────────────────────────
# "base" = good balance of speed/accuracy on 4GB VRAM
# Switch to "tiny" if RAM is tight during multi-file ingestion
_whisper_model = None

def _get_whisper():
    global _whisper_model
    if _whisper_model is None:
        print("[Whisper] Loading model (base)...")
        _whisper_model = whisper.load_model("base")
    return _whisper_model


# ── Per-format extractors ────────────────────────────────────────────────────

def extract_pdf(filepath: str) -> str:
    """Extract text from PDF. Falls back to OCR for scanned pages."""
    doc = fitz.open(filepath)
    parts = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()

        if len(text) > 50:
            # Digital PDF — native text available
            parts.append(f"[Page {page_num + 1}]\n{text}")
        else:
            # Scanned page — render to image and OCR
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img).strip()
            if ocr_text:
                parts.append(f"[Page {page_num + 1} — OCR]\n{ocr_text}")

    doc.close()
    return "\n\n".join(parts)


def extract_docx(filepath: str) -> str:
    """Extract text from DOCX including tables."""
    document = docx.Document(filepath)
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_image(filepath: str) -> str:
    """OCR an image file."""
    img = Image.open(filepath).convert("RGB")
    # Upscale small images for better OCR
    w, h = img.size
    if w < 1000:
        img = img.resize((w * 2, h * 2), Image.LANCZOS)
    text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
    return text.strip()


def extract_audio(filepath: str) -> str:
    """Transcribe audio using Whisper offline."""
    model = _get_whisper()
    result = model.transcribe(filepath, fp16=False)   # fp16=False → safer on mixed hardware
    return result["text"].strip()


# ── Format router ────────────────────────────────────────────────────────────

SUPPORTED_FORMATS = {
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".doc":  extract_docx,
    ".png":  extract_image,
    ".jpg":  extract_image,
    ".jpeg": extract_image,
    ".bmp":  extract_image,
    ".tiff": extract_image,
    ".tif":  extract_image,
    ".mp3":  extract_audio,
    ".mp4":  extract_audio,
    ".wav":  extract_audio,
    ".m4a":  extract_audio,
    ".ogg":  extract_audio,
}


def extract_text(filepath: str) -> str:
    """Route file to the correct extractor based on extension."""
    ext = Path(filepath).suffix.lower()
    extractor = SUPPORTED_FORMATS.get(ext)

    if extractor is None:
        raise ValueError(f"Unsupported file type: {ext}")

    print(f"[Ingest] Extracting from {Path(filepath).name} ({ext})")
    return extractor(filepath)


# ── Chunker ──────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 512, chunk_overlap: int = 64) -> list[str]:
    """
    Split text into overlapping chunks for better retrieval.
    chunk_size=512 tokens ≈ ~350–400 words — good for semantic search.
    chunk_overlap=64 prevents context loss at boundaries.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_text(text)


# ── Main pipeline entry ──────────────────────────────────────────────────────

def process_file(filepath: str) -> dict:
    """
    Full pipeline: file → text → chunks.
    Returns a dict with metadata + chunks ready for vector storage.
    """
    filename = Path(filepath).name
    ext = Path(filepath).suffix.lower()

    raw_text = extract_text(filepath)

    if not raw_text.strip():
        raise ValueError(f"No text could be extracted from {filename}")

    chunks = chunk_text(raw_text)

    print(f"[Ingest] {filename} → {len(raw_text)} chars → {len(chunks)} chunks")

    return {
        "filename": filename,
        "filepath": filepath,
        "file_type": ext.lstrip("."),
        "raw_text": raw_text,
        "chunks": chunks,
        "num_chunks": len(chunks),
    }
